from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import markdown

# ==========================
# Configuration
# ==========================
INPUT_MD = "Final_Project_Secure_Azure_Cloud_Report_Template.md"
OUTPUT_DOCX = "Final_Project_Secure_Azure_Cloud_Report_Template.docx"

# Convert Markdown to HTML
with open(INPUT_MD, "r", encoding="utf-8") as f:
    md_text = f.read()

html = markdown.markdown(
    md_text,
    extensions=["tables", "fenced_code"]
)

soup = BeautifulSoup(html, "html.parser")

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_table(html_table):
    rows = html_table.find_all("tr")
    cols = len(rows[0].find_all(["td", "th"]))

    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"

    for r in rows:
        cells = r.find_all(["td", "th"])
        row_cells = table.add_row().cells

        for i, cell in enumerate(cells):
            row_cells[i].text = cell.get_text(" ", strip=True)

    doc.add_paragraph()


for element in soup.children:

    if element.name is None:
        continue

    if element.name == "h1":
        doc.add_heading(element.get_text(), level=1)

    elif element.name == "h2":
        doc.add_heading(element.get_text(), level=2)

    elif element.name == "h3":
        doc.add_heading(element.get_text(), level=3)

    elif element.name == "h4":
        doc.add_heading(element.get_text(), level=4)

    elif element.name == "p":
        doc.add_paragraph(element.get_text())

    elif element.name == "ul":
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(), style="List Bullet")

    elif element.name == "ol":
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(), style="List Number")

    elif element.name == "blockquote":
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.italic = True

    elif element.name == "pre":
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.font.name = "Courier New"

    elif element.name == "table":
        add_table(element)

doc.save(OUTPUT_DOCX)

print(f"Successfully created: {OUTPUT_DOCX}")